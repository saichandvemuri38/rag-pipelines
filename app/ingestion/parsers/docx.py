import docx
from app.ingestion.formatter import Document

def load_docx(file_path: str) -> list[Document]:
    doc = docx.Document(file_path)
    full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    return [
        Document(
            text=full_text,
            metadata={
                "source": file_path,
                "type": "docx",
            },
        )
    ]